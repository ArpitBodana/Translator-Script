from docx import Document
from googletrans import Translator
from google_trans_new import google_translator  
from time import sleep

translator2 = Translator()
translator3 = Translator()
translator = google_translator()

paras=[]
check = False
checksen="Anchor script (Voice Over)"
endsen="Visual Reference"
'''Enter filename below  and put extension (.docx) at end '''
filename="your filename.docx"
input=Document(filename)
tables = input.tables
for table in tables:
    for column in table.columns:
        for cell in column.cells:
            mypara=cell.paragraphs
            i=0
            for paragraph in cell.paragraphs:
                    
                print(paragraph.text) 
                p=paragraph.text
                if p==endsen:
                    print('endpoint found!!')
                    print("Translation is completed please wait few seconds to save file !!")
                    check=False
                    sleep(3)
                    input.save(filename)
                    break
                if p==checksen:
                    check=True
                    print('startPoint found !!')
                    sleep(1)
                if check ==True:
                    
                    try:
                        translate_text=translator2.translate(p,dest="hi")
                        detect = translator.detect(translate_text)
                        # print(detect[0])
                        if(detect[0]=="en"):
                            raise ValueError("API Timeout!!")  
                        print(translate_text.text)            
                    except Exception as error:
                        # print(error)
                        # print('switching method for translation !!')
                        translate_text = translator.translate(p, lang_tgt='hi')
                        print(translate_text)    
                    finally:
                        try:
                            translate_text.text="\n"+translate_text.text
                            mypara[i].add_run(translate_text.text)
                            i=i+1
        
                        except Exception as err:
                            # print(err)
                            print("switching method to save!!")
                            sleep(1)
                            translate_text="\n"+translate_text
                            mypara[i].add_run(translate_text)
                            i=i+1
                
                            
                           



                                                            ##ArpitBodana2021                         

                

                        


